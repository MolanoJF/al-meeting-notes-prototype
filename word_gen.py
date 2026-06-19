"""
Generates a branded Ackroyd Lowrie Word document from meeting notes.

Accepts the sections-based dict produced by notion_utils.extract_meeting_summary():
    {
        "sections":     [{"heading": str, "lines": [str], "is_actions": bool}],
        "action_items": [{"who": str, "what": str, "by_when": str}],
    }

Styling matches AL brand: Arial throughout, logo-only header, A4.
"""

import os
import tempfile
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm, Emu

HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)
TITLE_COLOR   = RGBColor(0x00, 0x00, 0x00)
SUBTITLE_COLOR = RGBColor(0x6B, 0x72, 0x80)

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "al_logo.jpg")
LOGO_WIDTH_EMU  = 1238250
LOGO_HEIGHT_EMU = 314325

A4_WIDTH  = Mm(210)
A4_HEIGHT = Mm(297)


def _set_base_font(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Arial")


def _style_title(doc: Document):
    t = doc.styles["Title"]
    t.font.name = "Arial"
    t.font.size = Pt(28)
    t.font.bold = False
    t.font.color.rgb = TITLE_COLOR


def _style_heading1(doc: Document):
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(11)
    h1.font.bold = True
    h1.font.color.rgb = HEADING_COLOR


def _style_list_bullet(doc: Document):
    lb = doc.styles["List Bullet"]
    lb.font.name = "Arial"
    lb.font.size = Pt(10)


def _set_page_geometry(doc: Document):
    s = doc.sections[0]
    s.page_width    = A4_WIDTH
    s.page_height   = A4_HEIGHT
    s.top_margin    = Pt(54)
    s.left_margin   = Pt(54)
    s.right_margin  = Pt(54)
    s.bottom_margin = Pt(45)


def _add_header_logo(doc: Document):
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(LOGO_PATH):
        run = para.add_run()
        run.add_picture(LOGO_PATH, width=Emu(LOGO_WIDTH_EMU), height=Emu(LOGO_HEIGHT_EMU))


def _set_cell_shading(cell, hex_fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _add_action_table(doc: Document, action_items: list):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "WHO"
    hdr[1].text = "WHAT"
    hdr[2].text = "BY WHEN"
    for cell in hdr:
        _set_cell_shading(cell, "E5E7EB")   # light grey header
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            run.bold = True

    for item in action_items:
        row = table.add_row().cells
        row[0].text = item.get("who", "")
        row[1].text = item.get("what", "")
        row[2].text = item.get("by_when", "TBD")

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def generate_word_doc(summary_data: dict, meeting_meta: dict) -> str:
    """
    Generate an AL-branded Word doc.

    summary_data — output of notion_utils.extract_meeting_summary()
    meeting_meta — {"title": str, "date": str, "attendees": list[str]}
    """
    title     = meeting_meta.get("title", "Meeting Notes")
    date      = meeting_meta.get("date", datetime.today().strftime("%Y-%m-%d"))
    attendees = meeting_meta.get("attendees", [])

    sections     = summary_data.get("sections", [])
    action_items = summary_data.get("action_items", [])

    doc = Document()
    _set_base_font(doc)
    _style_title(doc)
    _style_heading1(doc)
    _style_list_bullet(doc)
    _set_page_geometry(doc)
    _add_header_logo(doc)

    # Cover: title + subtitle
    doc.add_heading(title, level=0)
    subtitle_parts = [date]
    if attendees:
        subtitle_parts.append(", ".join(attendees))
    subtitle = doc.add_paragraph("  ·  ".join(subtitle_parts))
    for run in subtitle.runs:
        run.font.color.rgb = SUBTITLE_COLOR

    # Render each section
    for section in sections:
        heading = section["heading"]
        lines   = section["lines"]

        doc.add_heading(heading, level=1)

        if section["is_actions"] and action_items:
            # Action items get a table (richer than a bullet list)
            _add_action_table(doc, action_items)
        else:
            for line in lines:
                doc.add_paragraph(line, style="List Bullet")

    # If action_items exist but no section claimed them, append a table at the end
    actions_rendered = any(s["is_actions"] for s in sections)
    if action_items and not actions_rendered:
        doc.add_heading("Action Items", level=1)
        _add_action_table(doc, action_items)

    tmp = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx",
        prefix=f"{date}-AL-",
    )
    doc.save(tmp.name)
    return tmp.name
