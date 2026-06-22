#!/usr/bin/env python3
"""
Generate AL Meeting Notes DOCX from template.

CLI usage:
    python scripts/generate_docx.py <input_json> <output_docx>

Importable usage:
    from scripts.generate_docx import generate_docx
    generate_docx(data_dict, Path("output.docx"))
"""

import copy
import json
import sys
from pathlib import Path

from docx import Document

TEMPLATE_PATH = Path(__file__).parent.parent / "assets" / "template.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_run_text(para, new_text):
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)


def replace_value_run(para, new_text):
    """For 'Label  [value]' paragraphs: replace all non-bold value runs with new_text."""
    value_runs = [r for r in para.runs if r.bold is not True]
    if not value_runs:
        para.add_run(new_text)
        return
    value_runs[0].text = new_text
    for r in value_runs[1:]:
        r.text = ""


def fill_left_column(cell, data):
    present = data.get("present", [])
    apologies = data.get("apologies", [])

    state = None
    practice_idx = 0
    name_idx = 0

    for para in cell.paragraphs:
        text = para.text.strip()

        if text == "PROJECT":
            state = "project"
            continue
        elif text == "SUBJECT":
            state = "subject"
            continue
        elif text == "DATE":
            state = "date"
            continue
        elif text == "REVISION":
            state = "revision"
            continue
        elif text == "PRESENT":
            state = "present"
            practice_idx = 0
            name_idx = 0
            continue
        elif text == "APOLOGIES":
            state = "apologies"
            continue

        if "[" not in text:
            continue

        if state == "project":
            set_run_text(para, data.get("project", ""))
            state = None
        elif state == "subject":
            set_run_text(para, data.get("meeting_type", ""))
            state = None
        elif state == "date":
            set_run_text(para, data.get("date", ""))
            state = None
        elif state == "revision":
            set_run_text(para, data.get("revision", "A"))
            state = None
        elif state == "present":
            if "[Practice name]" in text:
                name = present[practice_idx].get("practice", "") if practice_idx < len(present) else ""
                set_run_text(para, name)
                practice_idx += 1
                name_idx = 0
            elif "[Name" in text:
                p_idx = practice_idx - 1
                if 0 <= p_idx < len(present):
                    people = present[p_idx].get("people", [])
                    set_run_text(para, people[name_idx] if name_idx < len(people) else "")
                else:
                    set_run_text(para, "")
                name_idx += 1
        elif state == "apologies":
            if apologies:
                set_run_text(para, ", ".join(apologies))
            else:
                set_run_text(para, "-")
            state = None


def fill_right_column(cell, data):
    for para in cell.paragraphs:
        text = para.text
        if "[Project name]" in text and "Meeting Notes" not in text:
            set_run_text(para, data.get("project", ""))
        elif "Issued by" in text and "[" in text:
            replace_value_run(para, data.get("issued_by", ""))
        elif "Issue date" in text and "[" in text:
            replace_value_run(para, data.get("issue_date", ""))
        elif "Distribution" in text and "[" in text:
            replace_value_run(para, data.get("distribution", ""))
        elif "Next meeting" in text and "[" in text:
            replace_value_run(para, data.get("next_meeting", "TBD"))


def build_content_table(tbl, sections):
    """Returns list of action items: [{"ref": "X.Y", "comment": ..., "action": ...}]"""
    rows = list(tbl.rows)

    if len(rows) < 3:
        print("WARNING: Content table has fewer rows than expected.")
        return []

    section_header_template = copy.deepcopy(rows[1]._tr)
    item_template = copy.deepcopy(rows[2]._tr)

    tbl_element = tbl._tbl
    for row in rows[1:]:
        tbl_element.remove(row._tr)

    action_items = []
    content_num = 0

    for section in sections:
        heading = section.get("heading", "")
        items = section.get("items", [])
        is_action_section = heading.strip().lower() in ("action items", "actions")

        if not is_action_section:
            content_num += 1
            sh_tr = copy.deepcopy(section_header_template)
            t_els = sh_tr.findall(".//{%s}t" % W)
            if len(t_els) >= 1:
                t_els[0].text = f"{content_num}.0"
            if len(t_els) >= 2:
                t_els[1].text = heading
            tbl_element.append(sh_tr)

        for item_num, item in enumerate(items, start=1):
            ref = f"{content_num}.{item_num}"
            if not is_action_section:
                item_tr = copy.deepcopy(item_template)
                t_els = item_tr.findall(".//{%s}t" % W)
                if len(t_els) >= 1:
                    t_els[0].text = ref
                if len(t_els) >= 2:
                    t_els[1].text = item.get("comment", "")
                if len(t_els) >= 3:
                    t_els[2].text = item.get("action", "")
                tbl_element.append(item_tr)

            if item.get("action", "").strip():
                action_items.append({
                    "ref": ref,
                    "comment": item.get("comment", ""),
                    "action": item.get("action", ""),
                    "due": item.get("due", "-"),
                })

    return action_items


def build_name_to_practice(present):
    mapping = {}
    for group in present:
        practice = group.get("practice", "")
        for person in group.get("people", []):
            first_word = person.split()[0] if person else ""
            mapping[first_word.lower()] = practice
            name_part = person.split("—")[0].split("-")[0].strip()
            mapping[name_part.lower()] = practice
    return mapping


def build_action_register(tbl, action_items, name_to_practice):
    rows = list(tbl.rows)

    if len(rows) < 2:
        print("WARNING: Action register table has fewer rows than expected.")
        return

    row_template = copy.deepcopy(rows[1]._tr)
    tbl_element = tbl._tbl
    for row in rows[1:]:
        tbl_element.remove(row._tr)

    for item in action_items:
        owner_raw = item.get("action", "")
        first_word = owner_raw.split()[0].lower() if owner_raw else ""
        practice = name_to_practice.get(owner_raw.lower()) or name_to_practice.get(first_word) or owner_raw

        row_tr = copy.deepcopy(row_template)
        t_els = row_tr.findall(".//{%s}t" % W)
        values = [item["ref"], item["comment"], practice, item.get("due", "-"), "Open"]
        for i, val in enumerate(values):
            if i < len(t_els):
                t_els[i].text = val
        tbl_element.append(row_tr)

    if not action_items:
        row_tr = copy.deepcopy(row_template)
        tbl_element.append(row_tr)


def generate_docx(data: dict, output_path: Path) -> None:
    """Fill the AL template with meeting data and save to output_path."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))

    if len(doc.tables) < 2:
        raise ValueError(f"Expected ≥2 tables in template, found {len(doc.tables)}")

    fill_left_column(doc.tables[0].cell(0, 0), data)
    fill_right_column(doc.tables[0].cell(0, 1), data)

    action_items = build_content_table(doc.tables[1], data.get("sections", []))

    if len(doc.tables) >= 3:
        name_to_practice = build_name_to_practice(data.get("present", []))
        build_action_register(doc.tables[2], action_items, name_to_practice)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/generate_docx.py <input_json> <output_docx>")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"ERROR: Input JSON not found: {input_path}")
        sys.exit(1)

    data = json.loads(input_path.read_text(encoding="utf-8-sig"))
    generate_docx(data, output_path)
    print(f"[OK] Saved: {output_path}")


if __name__ == "__main__":
    main()
